import io
import streamlit as st
import pandas as pd
from datetime import date
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from model_inference import analyze_sentiment, analyze_batch
from utils import sentiment_visualizer as viz
from utils import data_handler as dh
from utils.utils import save_chart_as_image

st.set_page_config(page_title="E-Consultation Sentiment Analysis", layout="wide")

# --- SESSION STATE ---
if "comments" not in st.session_state:
    st.session_state["comments"] = []

# --- SIDEBAR ---
st.sidebar.title("➕ Add Comments")
input_method = st.sidebar.radio("Choose input method:", ["Single Comment", "Multiple Comments", "Upload File"])

if input_method == "Single Comment":
    text = st.sidebar.text_area("Comment text:")
    author = st.sidebar.text_input("Author (optional):")
    date_input = st.sidebar.date_input("Date", date.today())
    if st.sidebar.button("Add Comment"):
        if text.strip():
            st.session_state["comments"].append({"text": text, "author": author, "date": str(date_input)})
        else:
            st.sidebar.error("⚠️ Please enter a valid comment.")

elif input_method == "Multiple Comments":
    texts = st.sidebar.text_area("Enter multiple comments (one per line):")
    author = st.sidebar.text_input("Author (optional):")
    date_input = st.sidebar.date_input("Date", date.today())
    if st.sidebar.button("Add Comments"):
        if texts.strip():
            for t in texts.split("\n"):
                if t.strip():
                    st.session_state["comments"].append({"text": t.strip(), "author": author, "date": str(date_input)})
        else:
            st.sidebar.error("⚠️ Please enter at least one valid comment.")

elif input_method == "Upload File":
    uploaded_file = st.sidebar.file_uploader(
        "Upload CSV / XLSX / TXT with comments", 
        type=["csv", "txt", "xlsx", "xls"]
    )
    if uploaded_file is not None:
        try:
            df = dh.load_comments(uploaded_file)
            st.session_state["comments"].extend(df.to_dict("records"))
            st.success(f"✅ Loaded {len(df)} comments from file")
        except Exception as e:
            st.error(f"❌ Could not load file: {e}")

if st.sidebar.button("Clear All Comments"):
    st.session_state["comments"] = []

# --- EMPTY INPUT HANDLING ---
def display_empty_message():
    st.warning("⚠️ No comments to analyze. Please add comments via the sidebar.")

# --- MAIN APP ---
st.title("💬 E-Consultation Sentiment Analysis")
st.write("Analyze public sentiment from consultation comments with AI-powered insights")

tabs = st.tabs(["📊 Dashboard", "📈 Analytics", "💬 Comments View", "🔍 Insights", "📑 Reports"])

# --- Word Export Function ---
def export_to_word(df: pd.DataFrame, overall_stats: dict) -> bytes:
    doc = Document()
    doc.add_heading("E-Consultation Sentiment Analysis Report", 0)
    doc.add_heading("📊 Overall Statistics", level=1)
    for key, value in overall_stats.items():
        doc.add_paragraph(f"{key}: {value}")
    doc.add_heading("📝 Sample Results", level=1)
    table = doc.add_table(rows=1, cols=len(df.columns))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = col
    for _, row in df.head(20).iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

# --- PDF Export Function ---
def export_to_pdf(df: pd.DataFrame, overall_stats: dict, charts: list) -> bytes:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    elements = []
    elements.append(Paragraph("E-Consultation Sentiment Analysis Report", styles["Title"]))
    elements.append(Spacer(1, 12))
    elements.append(Paragraph("📊 Overall Statistics", styles["Heading2"]))
    for key, value in overall_stats.items():
        elements.append(Paragraph(f"{key}: {value}", styles["Normal"]))
    elements.append(Spacer(1, 12))
    elements.append(Paragraph("📝 Sample Results", styles["Heading2"]))
    table_data = [list(df.columns)] + df.head(15).values.tolist()
    t = Table(table_data)
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.black)
    ]))
    elements.append(t)
    for chart_path in charts:
        elements.append(Spacer(1, 20))
        elements.append(Image(chart_path, width=400, height=250))
    doc.build(elements)
    pdf = buffer.getvalue()
    buffer.close()
    return pdf

# TAB 1: Dashboard
with tabs[0]:
    st.subheader("Overview")
    if not st.session_state["comments"]:
        display_empty_message()
    else:
        if st.button("Analyze All Comments"):
            results = analyze_batch(st.session_state["comments"])
            st.session_state["results"] = results
            st.success("✅ Analysis completed!")
        if "results" in st.session_state:
            df = pd.DataFrame(st.session_state["results"])
            total = len(df)
            positive = len(df[df["sentiment_sub"] == "Positive"])
            negative = len(df[df["sentiment_sub"] == "Negative"])
            neutral_pure = len(df[df["sentiment_sub"] == "Neutral (Pure Neutral)"])
            neutral_neg = len(df[df["sentiment_sub"] == "Neutral (Dominantly Negative)"])
            neutral_pos = len(df[df["sentiment_sub"] == "Neutral (Dominantly Positive)"])
            col1, col2, col3, col4, col5, col6 = st.columns(6)
            col1.metric("😊 Positive", positive, f"{positive/total:.1%}")
            col2.metric("😞 Negative", negative, f"{negative/total:.1%}")
            col3.metric("😐 Neutral", neutral_pure, f"{neutral_pure/total:.1%}")
            col4.metric("😐⬇️ Neutral-Dom Neg", neutral_neg, f"{neutral_neg/total:.1%}")
            col5.metric("😐⬆️ Neutral-Dom Pos", neutral_pos, f"{neutral_pos/total:.1%}")
            col6.metric("📊 Total", total)
            c1, c2 = st.columns(2)
            fig_dist = viz.sentiment_distribution(df)
            fig_time = viz.sentiment_over_time(df)
            for fig in [fig_dist, fig_time]:
                fig.update_layout(
                    template="plotly_dark",
                    plot_bgcolor="black",
                    paper_bgcolor="black",
                    font=dict(color="white")
                )
            c1.plotly_chart(fig_dist, use_container_width=True)
            c2.plotly_chart(fig_time, use_container_width=True)
            c3 = st.columns(1)[0]
            fig_breakdown = viz.sentiment_main_sub_breakdown(df)
            c3.plotly_chart(fig_breakdown, use_container_width=True)

# TAB 2: Analytics
with tabs[1]:
    if not st.session_state["comments"]:
        display_empty_message()
    elif "results" in st.session_state:
        st.dataframe(pd.DataFrame(st.session_state["results"]).describe(include="all"))

# TAB 3: Comments View
with tabs[2]:
    if not st.session_state["comments"]:
        display_empty_message()
    elif "results" in st.session_state:
        st.dataframe(pd.DataFrame(st.session_state["results"]))

# TAB 4: Insights
with tabs[3]:
    if not st.session_state["comments"]:
        display_empty_message()
    else:
        st.info("✨ Future: Add word clouds, key phrases, etc.")

# TAB 5: Reports
with tabs[4]:
    if not st.session_state["comments"]:
        display_empty_message()
    elif "results" in st.session_state:
        df = pd.DataFrame(st.session_state["results"])
        expected_cols = ["text", "author", "date", "sentiment_main", "sentiment_sub", "score"]
        for col in expected_cols:
            if col not in df.columns:
                df[col] = ""
        df = df[expected_cols]
        st.dataframe(df)
        stats = {
            "Total Comments": len(df),
            "Positive": (df["sentiment_main"] == "Positive").sum(),
            "Negative": (df["sentiment_main"] == "Negative").sum(),
            "Neutral": (df["sentiment_main"] == "Neutral").sum(),
        }
        chart_paths = []
        fig_dist = viz.sentiment_distribution(df)
        fig_time = viz.sentiment_over_time(df)
        fig_breakdown = viz.sentiment_main_sub_breakdown(df)
        chart_paths.append(save_chart_as_image(fig_dist, "dist_chart.png"))
        chart_paths.append(save_chart_as_image(fig_time, "time_chart.png"))
        chart_paths.append(save_chart_as_image(fig_breakdown, "breakdown_chart.png"))
        chart_paths = [p for p in chart_paths if p]
        st.download_button(
            "📥 Download CSV",
            data=df.to_csv(index=False),
            file_name="sentiment_results.csv",
            mime="text/csv"
        )
        excel_buffer = io.BytesIO()
        with pd.ExcelWriter(excel_buffer, engine="openpyxl") as writer:
            df.to_excel(writer, index=False, sheet_name="Sentiment Results")
        st.download_button(
            "📥 Download Excel",
            data=excel_buffer.getvalue(),
            file_name="sentiment_results.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        word_bytes = export_to_word(df, stats)
        st.download_button(
            "📥 Download Word",
            data=word_bytes,
            file_name="sentiment_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        pdf_bytes = export_to_pdf(df, stats, charts=chart_paths)
        st.download_button(
            "📥 Download PDF",
            data=pdf_bytes,
            file_name="sentiment_report.pdf",
            mime="application/pdf"
        )
